"""DOCX export helpers for writing drafts."""
from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    DOCX_AVAILABLE = False
    Document = None


def _add_markdown_line(document: Document, line: str) -> None:
    if line.startswith("# "):
        document.add_heading(line[2:].strip(), level=1)
    elif line.startswith("## "):
        document.add_heading(line[3:].strip(), level=2)
    elif line.startswith("### "):
        document.add_heading(line[4:].strip(), level=3)
    elif line.startswith("- "):
        document.add_paragraph(line[2:].strip(), style="List Bullet")
    else:
        document.add_paragraph(line)


def build_docx_from_markdown(markdown: str, output_path: Path) -> Path:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Please add python-docx to requirements.")

    document = Document()
    for line in markdown.splitlines():
        if not line.strip():
            continue
        _add_markdown_line(document, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
